from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def guardar_texto_en_word(nombre_archivo):
    # Crear un documento Word
    doc = Document()
    
    # Estilo personalizado para el encabezado principal
    estilo_titulo = doc.styles['Heading 1']
    estilo_titulo.font.size = Pt(16)
    
    # Estilo para el encabezado de nivel 2
    estilo_subtitulo = doc.styles['Heading 2']
    estilo_subtitulo.font.size = Pt(14)
    
    while True:
        # Solicitar al usuario que ingrese el texto completo
        texto_caso = input("Ingresa el caso de uso completo (o 'salir' para terminar): ")
        if texto_caso.lower() == 'salir':
            break
        
        # Separar las secciones por líneas y luego por etiquetas
        lineas = texto_caso.split('\n')
        caso_uso = lineas[0].strip()
        doc.add_heading(caso_uso, level=1)  # Encabezado del caso de uso
        
        # Diccionario de secciones
        secciones = {
            'Objetivos asociados:': 'Objetivos asociados',
            'Actor/es Relacionados:': 'Actor/es Relacionados',
            'Requisitos asociados:': 'Requisitos asociados',
            'Descripción:': 'Descripción',
            'Precondición:': 'Precondición',
            'Secuencia Normal:': 'Secuencia Normal',
            'Post-condición:': 'Post-condición',
            'Excepciones:': 'Excepciones',
            'Rendimiento:': 'Rendimiento',
            'Frecuencia esperada:': 'Frecuencia esperada',
            'Comentarios:': 'Comentarios',
        }
        
        current_section = None
        
        for linea in lineas[1:]:
            linea = linea.strip()
            if linea in secciones:
                # Crear nueva tabla con dos columnas (Título, Contenido)
                tabla = doc.add_table(rows=1, cols=2)
                tabla.style = 'Table Grid'
                hdr_cells = tabla.rows[0].cells
                hdr_cells[0].text = secciones[linea]  # Columna de Título
                hdr_cells[0].paragraphs[0].runs[0].bold = True
                hdr_cells[0].width = Pt(100)  # Ajustar el ancho de la columna de título
                
                current_section = secciones[linea]
                
            elif current_section:
                # Colocar el contenido en la segunda columna de la tabla
                fila = tabla.add_row().cells
                fila[0].text = ""  # Dejar la primera celda en blanco para la alineación de la tabla
                fila[1].text = linea  # Agregar el contenido en la segunda celda
                
                if current_section in ['Secuencia Normal', 'Excepciones', 'Rendimiento']:
                    # Agregar pasos a la tabla cuando sea necesario
                    pasos = linea.split(';')  # Se asume que los pasos están separados por ';'
                    for paso in pasos:
                        fila = tabla.add_row().cells
                        datos_paso = paso.split(':')  # Separar número de paso y descripción
                        fila[0].text = datos_paso[0] if len(datos_paso) > 0 else ''
                        fila[1].text = datos_paso[1] if len(datos_paso) > 1 else ''
    
    # Guardar el documento
    doc.save(nombre_archivo)
    print(f"El documento se ha guardado en {nombre_archivo}")

def main():
    nombre_archivo = "documento_casos_uso.docx"  # Nombre del archivo de Word
    guardar_texto_en_word(nombre_archivo)

if __name__ == "__main__":
    main()
